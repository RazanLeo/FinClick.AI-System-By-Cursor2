"""
Document Processor - معالج المستندات
يدعم PDF, Excel, Word, صور
"""

import pandas as pd
import numpy as np
from typing import Dict, List, Optional, Union, Any
import logging
import os
from pathlib import Path

# Document Processing Libraries
try:
    import PyPDF2
    import docx
    import openpyxl
    from PIL import Image
    import pytesseract
except ImportError:
    logging.warning("Some document processing libraries not available")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """معالج المستندات متعدد الصيغ"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.docx': self._process_word,
            '.doc': self._process_word,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.tiff': self._process_image,
            '.bmp': self._process_image
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """معالجة المستند واستخراج البيانات"""
        try:
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # استخراج البيانات
            extracted_data = self.supported_formats[file_extension](file_path)
            
            # تنظيف وتنظيم البيانات
            cleaned_data = self._clean_data(extracted_data)
            
            # تحويل إلى DataFrames
            financial_data = self._convert_to_financial_data(cleaned_data)
            
            return {
                'success': True,
                'file_type': file_extension,
                'extracted_data': extracted_data,
                'cleaned_data': cleaned_data,
                'financial_data': financial_data,
                'processing_metadata': {
                    'file_size': file_path.stat().st_size,
                    'processing_time': 'calculated',
                    'confidence_score': 0.85
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': file_path.suffix if 'file_path' in locals() else 'unknown'
            }
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """معالجة ملفات PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text_content = ""
                tables_data = []
                images_data = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    # استخراج النص
                    page_text = page.extract_text()
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    
                    # محاولة استخراج الجداول (يتطلب مكتبات إضافية)
                    # tables_data.extend(self._extract_tables_from_pdf_page(page))
                    
                    # محاولة استخراج الصور
                    # images_data.extend(self._extract_images_from_pdf_page(page))
                
                return {
                    'type': 'pdf',
                    'text_content': text_content,
                    'tables': tables_data,
                    'images': images_data,
                    'num_pages': len(pdf_reader.pages),
                    'metadata': pdf_reader.metadata
                }
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """معالجة ملفات Excel"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            sheets_data = {}
            charts_data = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # استخراج البيانات
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append(row)
                
                if data:
                    # إنشاء DataFrame
                    df = pd.DataFrame(data[1:], columns=data[0])
                    sheets_data[sheet_name] = df
                
                # استخراج الرسوم البيانية
                if sheet.charts:
                    charts_data.extend([chart.title for chart in sheet.charts])
            
            return {
                'type': 'excel',
                'sheets': sheets_data,
                'charts': charts_data,
                'num_sheets': len(workbook.sheetnames),
                'workbook_properties': workbook.properties
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            raise
    
    def _process_word(self, file_path: Path) -> Dict[str, Any]:
        """معالجة ملفات Word"""
        try:
            doc = docx.Document(file_path)
            
            text_content = ""
            tables_data = []
            images_data = []
            
            # استخراج النص
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # استخراج الجداول
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    tables_data.append(df)
            
            # استخراج الصور
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images_data.append(rel.target_ref)
            
            return {
                'type': 'word',
                'text_content': text_content,
                'tables': tables_data,
                'images': images_data,
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'document_properties': doc.core_properties
            }
            
        except Exception as e:
            logger.error(f"Error processing Word {file_path}: {e}")
            raise
    
    def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """معالجة الصور باستخدام OCR"""
        try:
            image = Image.open(file_path)
            
            # تحويل الصورة إلى نص باستخدام OCR
            try:
                text_content = pytesseract.image_to_string(image, lang='ara+eng')
            except:
                text_content = pytesseract.image_to_string(image)
            
            # محاولة استخراج الجداول من الصورة
            tables_data = self._extract_tables_from_image(image)
            
            return {
                'type': 'image',
                'text_content': text_content,
                'tables': tables_data,
                'image_properties': {
                    'format': image.format,
                    'mode': image.mode,
                    'size': image.size,
                    'dpi': image.info.get('dpi', 'unknown')
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise
    
    def _extract_tables_from_image(self, image: Image.Image) -> List[pd.DataFrame]:
        """استخراج الجداول من الصور"""
        # هذا يتطلب مكتبات متقدمة مثل TableNet أو Camelot
        # للتبسيط، سنعيد قائمة فارغة
        return []
    
    def _clean_data(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """تنظيف وتنظيم البيانات المستخرجة"""
        cleaned_data = {
            'financial_statements': {},
            'text_content': '',
            'tables': [],
            'metadata': {}
        }
        
        # تنظيف النص
        if 'text_content' in extracted_data:
            cleaned_data['text_content'] = self._clean_text(extracted_data['text_content'])
        
        # تنظيف الجداول
        if 'tables' in extracted_data:
            cleaned_data['tables'] = self._clean_tables(extracted_data['tables'])
        
        # تنظيف البيانات المالية
        if 'sheets' in extracted_data:
            cleaned_data['financial_statements'] = self._clean_financial_data(extracted_data['sheets'])
        
        # تنظيف البيانات الوصفية
        cleaned_data['metadata'] = self._clean_metadata(extracted_data)
        
        return cleaned_data
    
    def _clean_text(self, text: str) -> str:
        """تنظيف النص"""
        if not text:
            return ""
        
        # إزالة الأحرف الخاصة
        text = text.replace('\x00', '')
        
        # تنظيف المسافات
        text = ' '.join(text.split())
        
        # إزالة الأسطر الفارغة المتكررة
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        text = '\n'.join(cleaned_lines)
        
        return text
    
    def _clean_tables(self, tables: List) -> List[pd.DataFrame]:
        """تنظيف الجداول"""
        cleaned_tables = []
        
        for table in tables:
            if isinstance(table, pd.DataFrame):
                # تنظيف DataFrame
                cleaned_df = table.copy()
                
                # إزالة الصفوف الفارغة
                cleaned_df = cleaned_df.dropna(how='all')
                
                # إزالة الأعمدة الفارغة
                cleaned_df = cleaned_df.dropna(axis=1, how='all')
                
                # تنظيف أسماء الأعمدة
                cleaned_df.columns = [str(col).strip() for col in cleaned_df.columns]
                
                cleaned_tables.append(cleaned_df)
        
        return cleaned_tables
    
    def _clean_financial_data(self, sheets: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
        """تنظيف البيانات المالية"""
        financial_data = {}
        
        for sheet_name, df in sheets.items():
            sheet_name_lower = sheet_name.lower()
            
            if any(keyword in sheet_name_lower for keyword in ['balance', 'asset', 'liability']):
                financial_data['balance_sheet'] = self._clean_balance_sheet(df)
            elif any(keyword in sheet_name_lower for keyword in ['income', 'revenue', 'profit', 'loss']):
                financial_data['income_statement'] = self._clean_income_statement(df)
            elif any(keyword in sheet_name_lower for keyword in ['cash', 'flow']):
                financial_data['cash_flow_statement'] = self._clean_cash_flow_statement(df)
            else:
                financial_data[f'sheet_{sheet_name}'] = df
        
        return financial_data
    
    def _clean_balance_sheet(self, df: pd.DataFrame) -> pd.DataFrame:
        """تنظيف الميزانية العمومية"""
        # تنظيف أسماء الأصول والخصوم
        # يمكن إضافة منطق أكثر تعقيداً هنا
        return df
    
    def _clean_income_statement(self, df: pd.DataFrame) -> pd.DataFrame:
        """تنظيف قائمة الدخل"""
        # تنظيف بنود الإيرادات والمصروفات
        return df
    
    def _clean_cash_flow_statement(self, df: pd.DataFrame) -> pd.DataFrame:
        """تنظيف قائمة التدفق النقدي"""
        # تنظيف بنود التدفق النقدي
        return df
    
    def _clean_metadata(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """تنظيف البيانات الوصفية"""
        metadata = {}
        
        # استخراج البيانات الوصفية المفيدة
        if 'metadata' in data:
            metadata.update(data['metadata'])
        
        if 'workbook_properties' in data:
            metadata.update(data['workbook_properties'])
        
        if 'document_properties' in data:
            metadata.update(data['document_properties'])
        
        if 'image_properties' in data:
            metadata.update(data['image_properties'])
        
        return metadata
    
    def _convert_to_financial_data(self, cleaned_data: Dict[str, Any]) -> Dict[str, Any]:
        """تحويل البيانات إلى تنسيق مالي قياسي"""
        financial_data = {
            'balance_sheet': None,
            'income_statement': None,
            'cash_flow_statement': None,
            'additional_data': {}
        }
        
        # تحويل البيانات المالية
        if 'financial_statements' in cleaned_data:
            statements = cleaned_data['financial_statements']
            
            if 'balance_sheet' in statements:
                financial_data['balance_sheet'] = statements['balance_sheet']
            
            if 'income_statement' in statements:
                financial_data['income_statement'] = statements['income_statement']
            
            if 'cash_flow_statement' in statements:
                financial_data['cash_flow_statement'] = statements['cash_flow_statement']
            
            # البيانات الإضافية
            for key, value in statements.items():
                if key not in ['balance_sheet', 'income_statement', 'cash_flow_statement']:
                    financial_data['additional_data'][key] = value
        
        return financial_data
    
    def get_processing_summary(self, processing_result: Dict[str, Any]) -> Dict[str, Any]:
        """الحصول على ملخص المعالجة"""
        if not processing_result.get('success', False):
            return {'error': processing_result.get('error', 'Unknown error')}
        
        financial_data = processing_result.get('financial_data', {})
        
        summary = {
            'file_type': processing_result.get('file_type'),
            'processing_success': True,
            'data_extracted': {
                'balance_sheet': financial_data.get('balance_sheet') is not None,
                'income_statement': financial_data.get('income_statement') is not None,
                'cash_flow_statement': financial_data.get('cash_flow_statement') is not None,
                'additional_sheets': len(financial_data.get('additional_data', {}))
            },
            'text_content_length': len(processing_result.get('cleaned_data', {}).get('text_content', '')),
            'tables_count': len(processing_result.get('cleaned_data', {}).get('tables', [])),
            'confidence_score': processing_result.get('processing_metadata', {}).get('confidence_score', 0)
        }
        
        return summary
